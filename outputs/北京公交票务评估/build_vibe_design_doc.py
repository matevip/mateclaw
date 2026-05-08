from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path
from re import sub
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/mate/Docs/Work/北京公交集团票务综合管理平台系统升级-需求分析说明书（初稿0506）.docx")
OUT = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/北京公交集团票务综合管理平台系统升级-SDD实施方案与工作量清单.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(96, 108, 122)
LIGHT = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"


def clean(text: str) -> str:
    text = text.replace("\u200c", "").replace("\u200b", "").replace("\ufeff", "")
    text = sub(r"\s+", " ", text).strip()
    return text


def module_name(text: str) -> str:
    text = clean(text)
    text = sub(r"^[\d.、\s]+", "", text)
    return text.strip("‌ ")


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def table_width(table, widths: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(cell)


def add_para(doc, text="", bold=False, color=None, size=10.5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)


def bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=10)


def numbered(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=10)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], fill=LIGHT, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade(c, fill)
        margins(c)
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.02
                for r in p.runs:
                    set_run_font(r, size=font_size)
    table_width(table, widths)
    doc.add_paragraph()


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    table_width(table, [6.5])
    doc.add_paragraph()


def configure(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.text = "北京公交集团票务综合管理平台系统升级 | SDD 实施与工作量评估"
    for r in header.runs:
        set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "内部评估文件"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in footer.runs:
        set_run_font(r, size=8.5, color=MUTED)


def extract_items():
    doc = Document(SOURCE)
    table = doc.tables[38]
    items = []
    current = ""
    for row in table.rows[1:]:
        cells = [clean(c.text) for c in row.cells[:5]]
        mod, page, func, change, note = cells
        if mod:
            current = module_name(mod)
        if not page:
            continue
        items.append({
            "module": current,
            "page": page,
            "function": func,
            "change": change,
            "note": note,
        })
    return items


DELETE_OR_MERGE_PAGES = {"模版下载", "票单打印", "包专管理", "无人售管理（车辆）"}
GROUP_COMPANY_REPORTS = {
    "配出统计",
    "存售日报表",
    "普票收入日报表",
    "票款清分日报",
    "IC卡充值日报表",
    "IC卡充值走势报表",
    "收入走势报表",
    "大额币走势报表",
    "残币入账走势报表",
    "分公司票务管理月报表",
}


def normalize_requirement_items(items):
    normalized = []
    for item in items:
        item = dict(item)
        if item["page"] in DELETE_OR_MERGE_PAGES:
            item["change"] = "删除/整合"
            item["note"] = f"{item['note']} 原需求标注删除或合并，按菜单下线/能力整合处理。".strip()
        if item["module"] == "统计报表（分公司）" and item["page"] in GROUP_COMPANY_REPORTS:
            continue
        if item["module"] == "统计报表（集团）" and item["page"] in GROUP_COMPANY_REPORTS:
            item["module"] = "统计报表（集团/分公司整合）"
            item["note"] = f"{item['note']} 集团/分公司同名报表合并建设，按数据权限切换展示范围。".strip()
        normalized.append(item)
    return normalized


def add_external_and_pending_items(items):
    additions = [
        {
            "module": "外部系统对接",
            "page": "公交集团数据湖标准数据接口",
            "function": "统一数据接口、数据格式、传输频率、安全保障、数据加密",
            "change": "待补充接口文档、数据范围、推拉模式、频率和鉴权要求",
            "note": "需求文档 3.3.1 明确用于数据湖获取票务系统数据。",
        },
        {
            "module": "外部系统对接",
            "page": "二维码系统数据接口",
            "function": "统一标准数据接口、二维码相关历史/统计数据供给",
            "change": "需求文档中现有二维码报表拟删除，但对外接口仍需确认是否保留数据输出",
            "note": "需确认二维码普票进出站、售检统计是否仅删除页面，还是同步取消数据服务。",
        },
        {
            "module": "外部系统对接",
            "page": "银行清分清点系统线下对接",
            "function": "线下报表传递、手动导入、清分日报入库",
            "change": "银行无网环境运行，需通过文件模板、导入校验和人工确认完成",
            "note": "涉及无人售票与票款清分，需失败明细、重复导入控制和导入日志。",
        },
        {
            "module": "外部系统对接",
            "page": "票款清分/胆款系统同步",
            "function": "票务系统导入数据后同步胆款系统，涉及票款清分日报",
            "change": "拓展部分待明确，需确认同步字段、触发时点、失败重试和对账机制",
            "note": "建议作为独立联调里程碑，未提供接口前只做适配层和模拟接口。",
        },
        {
            "module": "外部系统对接",
            "page": "IC卡/一卡通系统数据对接",
            "function": "线路、车辆、司售、IC卡类型、充值相关数据同步",
            "change": "需确认接口方式、同步频率、主数据归属、异常补偿",
            "note": "需求文档多处说明线路、车辆、司售信息来自 IC 卡系统接口同步。",
        },
        {
            "module": "外部系统对接",
            "page": "电子签章系统预留/对接",
            "function": "特殊业务在线签字确认、签章状态、签章文件归档",
            "change": "根据业务后续评估再升级",
            "note": "一期建议预留字段和附件能力；正式联调、CA/签章授权另列。",
        },
        {
            "module": "待定事项",
            "page": "安全性与合规要求确认",
            "function": "安全策略、审计要求、加密要求、等保/密评/信创测评配合",
            "change": "需求文档标注待集团补充确认",
            "note": "会影响部署架构、日志留存、账号策略、接口加密和报价边界。",
        },
        {
            "module": "待定事项",
            "page": "调账管理业务场景确认",
            "function": "调账申请、审核、库存/金额调整、差异流水",
            "change": "具体业务场景待明确",
            "note": "需明确可调对象、审批层级、是否影响历史报表和会计口径。",
        },
        {
            "module": "待定事项",
            "page": "特殊票管理业务场景确认",
            "function": "特殊票新增、修改、删除、领取、核销或结算关系",
            "change": "具体业务场景待明确",
            "note": "需确认是否进入库存流水、是否纳入收入和报表。",
        },
        {
            "module": "待定事项",
            "page": "无人售票自动对账与数据清洗",
            "function": "无人售票清分数据自动对账、异常识别、数据清洗",
            "change": "文档提出新增能力，但规则和数据来源需补充",
            "note": "需确认银行文件格式、对账容差、异常处理流程和人工复核机制。",
        },
        {
            "module": "专项工作",
            "page": "数据迁移与上线割接",
            "function": "Oracle 到达梦迁移、三轮迁移演练、增量割接、抽样核验、回滚预案",
            "change": "专项工作量，不并入普通功能页面",
            "note": "原需求明确 79 张表、序列、视图、函数、约 200GB 历史数据和 10 年存储。",
        },
        {
            "module": "专项工作",
            "page": "培训与运维交接",
            "function": "管理员培训、业务培训、报表核对培训、运维培训、用户手册和培训课件",
            "change": "专项工作量",
            "note": "建议至少上线前集中培训和上线后复盘培训两轮。",
        },
        {
            "module": "平台底座专项",
            "page": "消息中心与审批待办",
            "function": "阈值预警、申请/审核消息、未读已读、弹窗提醒、待办跳转",
            "change": "原需求明确消息中心和阈值预警",
            "note": "作为平台底座能力单列，不混入单个业务模块。",
        },
        {
            "module": "平台底座专项",
            "page": "打印模板与浏览器打印",
            "function": "配票单、结算单、打印模板、纸张规格、浏览器打印兼容",
            "change": "票单打印菜单删除，但打印能力保留",
            "note": "打印能力整合至配票、结算、查询入口。",
        },
        {
            "module": "平台底座专项",
            "page": "Excel异步导出中心",
            "function": "大数据量导出、异步任务、下载中心、导出权限、水印/字段控制",
            "change": "报表和查询列表共用能力",
            "note": "避免大文件导出阻塞浏览器请求。",
        },
    ]
    return items + additions


def classify(item):
    text = f"{item['module']} {item['page']} {item['function']} {item['change']} {item['note']}"
    if item["module"] == "专项工作":
        if "数据迁移" in text:
            return "专项工作", 50, 80, "三轮迁移演练、数据治理、增量割接和回滚验证不可省略。"
        if "培训" in text:
            return "专项工作", 15, 25, "培训材料、分角色培训和上线后复盘需单列交付。"
    if item["module"] == "平台底座专项":
        if "消息" in text:
            return "平台能力", 8, 14, "消息中心、待办提醒和阈值预警属于跨模块底座能力。"
        if "打印" in text:
            return "平台能力", 8, 14, "打印菜单删除不代表打印能力删除，需统一模板和兼容性验证。"
        if "导出" in text:
            return "平台能力", 8, 14, "大数据量导出需异步任务、权限控制和下载中心。"
    if item["module"] == "外部系统对接":
        if "数据湖" in text:
            return "外部对接", 18, 35, "需确认数据范围、接口协议、鉴权加密、传输频率、推拉模式和验收样例。"
        if "胆款" in text or "清分" in text:
            return "外部对接", 15, 32, "需确认字段口径、文件/接口格式、重试补偿、对账机制和失败处理。"
        if "IC卡" in text or "一卡通" in text:
            return "外部对接", 18, 36, "需确认主数据归属、同步频率、增量机制、异常补偿和历史数据映射。"
        if "电子签章" in text:
            return "外部对接", 8, 20, "一期建议预留签章对象、签章状态和附件归档；正式联调视接口材料另估。"
        return "外部对接", 12, 28, "需以接口文档、联调环境、样例数据和验收规则为准。"
    if item["module"] == "待定事项":
        if "无人售票自动对账" in text:
            return "待定补充", 12, 24, "自动对账规则未冻结，需按银行清分护栏预留分析、校验和复核工时。"
        if "调账管理" in text:
            return "待定补充", 8, 24, "调账业务会影响库存、金额、报表和审计追溯，需高于普通待定事项估算。"
        return "待定补充", 6, 24, "当前需求材料不足，需业务或集团侧补充确认后才能冻结范围。"
    if "可删除" in text or "删除" in item["page"] or "删除" in item["change"]:
        return "拟删除/整合", 0, 1, "确认下线、菜单隐藏、历史数据保留和跳转清理。"
    if any(k in text for k in ["结算", "重新结算", "配票", "调票", "调账", "退票", "库存", "入库", "出库", "核销", "调拨", "印制", "票号"]):
        lo, hi = 8, 18
        if any(k in text for k in ["票单结算", "重新结算", "票袋调票", "票柜退票", "票袋退票"]):
            lo, hi = 14, 26
        return "核心业务", lo, hi, "需要状态机、库存流水、票号段校验、事务一致性和回归测试。"
    if any(k in item["module"] for k in ["统计报表"]) or "报表" in text or "统计" in text:
        return "报表统计", 6, 14, "需确认统计口径、权限范围、导出格式和历史数据性能。"
    if "导入" in text or "上传" in text or "下载" in text:
        return "导入导出", 5, 10, "需模板、字段校验、失败明细、导入日志和重复处理。"
    if any(k in item["module"] for k in ["系统管理"]) or any(k in item["page"] for k in ["用户", "权限", "菜单", "日志", "阈值"]):
        return "平台能力", 6, 14, "需统一认证授权、按钮权限、数据范围和审计记录。"
    if any(k in text for k in ["查询、新建、修改、删除", "添加、编辑、删除", "添加、编辑", "新增", "管理"]):
        return "常规功能", 4, 8, "以列表、表单、校验、权限按钮和操作日志为主。"
    return "查询维护", 3, 6, "以查询、筛选、详情、导出或基础维护为主。"


def build_rows(items):
    rows = []
    for idx, item in enumerate(items, start=1):
        kind, lo, hi, detail = classify(item)
        vibe_lo = round(lo * 0.32, 1)
        vibe_hi = round(hi * 0.40, 1)
        vibe_lo, vibe_hi = apply_protection_floor(item, kind, vibe_lo, vibe_hi)
        vibe_display = build_vibe_display(item, vibe_lo, vibe_hi)
        dev = build_dev_scope(item, kind)
        rows.append({
            **item,
            "idx": idx,
            "kind": kind,
            "lo": lo,
            "hi": hi,
            "vibe_lo": vibe_lo,
            "vibe_hi": vibe_hi,
            "vibe_display": vibe_display,
            "dev": dev,
            "detail": detail,
        })
    return rows


def build_vibe_display(item, vibe_lo, vibe_hi):
    text = f"{item['module']} {item['page']} {item['function']} {item['change']} {item['note']}"
    if "二维码系统数据接口" in text:
        return "0-1 / 12-18"
    return f"{vibe_lo}-{vibe_hi}"


def apply_protection_floor(item, kind, vibe_lo, vibe_hi):
    text = f"{item['module']} {item['page']} {item['function']} {item['change']} {item['note']}"
    if kind == "拟删除/整合":
        return 0, 1
    if item["module"] == "专项工作" and "数据迁移" in text:
        return 30, 55
    if item["module"] == "专项工作" and "培训" in text:
        return 10, 18
    if any(k in text for k in ["票单结算", "票单重新结算", "重新结算", "月终结算", "票袋退票", "票柜退票", "票袋调票", "调账"]):
        vibe_lo = max(vibe_lo, 8)
        vibe_hi = max(vibe_hi, 12)
    if any(k in text for k in ["数据湖", "IC卡/一卡通", "IC 卡/一卡通"]):
        vibe_lo = max(vibe_lo, 15)
        vibe_hi = max(vibe_hi, 20)
    if "胆款系统" in text:
        vibe_lo = max(vibe_lo, 15)
        vibe_hi = max(vibe_hi, 18)
    if "二维码系统数据接口" in text:
        vibe_lo = max(vibe_lo, 12)
        vibe_hi = max(vibe_hi, 18)
    if any(k in text for k in ["银行清分", "自动对账"]):
        vibe_lo = max(vibe_lo, 12)
        vibe_hi = max(vibe_hi, 18)
    if any(k in text for k in ["票款实交银行登记", "票务室交班对账"]):
        vibe_lo = max(vibe_lo, 6)
        vibe_hi = max(vibe_hi, 12)
    if "配票应急导入" in text:
        vibe_lo = max(vibe_lo, 6)
        vibe_hi = max(vibe_hi, 10)
    if kind == "报表统计":
        vibe_lo = max(vibe_lo, 4)
        vibe_hi = max(vibe_hi, 8)
    if vibe_hi < vibe_lo:
        vibe_hi = vibe_lo
    return round(vibe_lo, 1), round(vibe_hi, 1)


def build_dev_scope(item, kind):
    page = item["page"]
    if kind == "拟删除/整合":
        return f"确认 {page} 下线或合并路径，处理菜单、权限、历史查询入口和数据保留。"
    if kind == "核心业务":
        return f"建设 {page} 的单据、审批/确认、库存流水、票号段校验、查询导出和异常回滚。"
    if kind == "报表统计":
        return f"建设 {page} 查询、统计聚合、权限过滤、Excel 导出、必要趋势展示和口径校验。"
    if kind == "导入导出":
        return f"建设 {page} 模板下载、文件上传、字段校验、入库处理、失败明细和导入日志。"
    if kind == "平台能力":
        return f"建设 {page} 的配置管理、权限控制、审计记录、查询筛选和维护页面。"
    return f"建设 {page} 的列表查询、详情、新增/编辑/停用或删除、权限按钮、日志和导出能力。"


def module_summary(rows):
    result = []
    grouped = defaultdict(list)
    for row in rows:
        grouped[row["module"]].append(row)
    for mod, vals in grouped.items():
        lo = sum(v["lo"] for v in vals)
        hi = sum(v["hi"] for v in vals)
        vlo = sum(v["vibe_lo"] for v in vals)
        vhi = sum(v["vibe_hi"] for v in vals)
        focus = module_focus(mod)
        if mod == "外部系统对接":
            focus = "含数据湖、二维码、银行清分、胆款、IC 卡、电子签章 6 项。二维码若业务下线，模块协同后下调至 60-84 人日；若仍需对外供数，维持 72-102 人日。开工前第 1 周以集团业务部书面冻结结果为准。"
        result.append([mod, str(len(vals)), f"{lo}-{hi}", f"{round(vlo)}-{round(vhi)}", focus])
    report_vals = []
    for mod, vals in grouped.items():
        if "统计报表" in mod:
            report_vals.extend(vals)
    if report_vals:
        lo = sum(v["lo"] for v in report_vals)
        hi = sum(v["hi"] for v in report_vals)
        vlo = sum(v["vibe_lo"] for v in report_vals)
        vhi = sum(v["vibe_hi"] for v in report_vals)
        deleted = sum(1 for v in report_vals if v["kind"] == "拟删除/整合")
        actual = len(report_vals) - deleted
        result.append([
            "报表合计（含整合/集团/分公司/票务室）",
            f"{len(report_vals)}",
            f"{lo}-{hi}",
            f"{round(vlo)}-{round(vhi)}",
            f"报表实际开发 {actual} 项，拟删除/整合 {deleted} 项，便于报价沟通。",
        ])
    return result


def module_focus(mod):
    if "专项工作" in mod:
        return "数据迁移、上线割接、培训和运维交接专项，不应混入普通功能项。"
    if "平台底座专项" in mod:
        return "消息中心、打印模板、异步导出等跨模块共用能力。"
    if "外部系统" in mod:
        return "数据湖、二维码、银行清分、胆款、IC卡、电子签章等接口联调。"
    if "待定事项" in mod:
        return "需求边界未冻结，需补材料后确定开发、测试和报价口径。"
    if "票单" in mod:
        return "票袋、配票、结算、重新结算、票号追溯，是核心复杂域。"
    if "票柜" in mod or "票库" in mod or "申请" in mod or "审核" in mod:
        return "围绕库存流转、审批状态和票号段控制。"
    if "统计报表" in mod:
        return "以口径确认、聚合性能、导出格式和权限过滤为重点。"
    if "系统" in mod:
        return "全局权限、菜单、日志、阈值，是平台底座。"
    if "基础" in mod:
        return "主数据维护和外部同步，是后续业务依赖。"
    return "按业务场景完成查询、维护、导入导出和日志审计。"


def make_landscape(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    return section


def build():
    items = add_external_and_pending_items(normalize_requirement_items(extract_items()))
    rows = build_rows(items)
    normal_lo = sum(r["lo"] for r in rows)
    normal_hi = sum(r["hi"] for r in rows)
    vibe_lo = round(sum(r["vibe_lo"] for r in rows))
    vibe_hi = round(sum(r["vibe_hi"] for r in rows))

    doc = Document()
    configure(doc)

    add_para(doc, "项目实施方案", bold=True, color=MUTED, size=11)
    p = doc.add_paragraph()
    r = p.add_run("北京公交集团票务综合管理平台系统升级")
    set_run_font(r, size=23, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    r = p.add_run("SDD 实施方案与工作量清单")
    set_run_font(r, size=15, color=MUTED)
    p.paragraph_format.space_after = Pt(14)
    add_table(doc, ["字段", "内容"], [
        ["文档用途", "用于项目立项、排期、报价拆分、研发组织和验收范围确认"],
        ["文档版本", "v1.0"],
        ["编制日期", date.today().isoformat()],
        ["编制人", "项目经理 / 解决方案团队"],
        ["审核人", "架构师 / 技术负责人"],
        ["变更记录", "v0.9 初稿；v1.0 修订工作量护栏、删除项口径、外部接口区间、报表整合、部署运维披露、目标环境矩阵和 SDD 术语展开"],
        ["需求来源", "北京公交集团票务综合管理平台系统升级需求分析说明书（初稿0506）"],
        ["估算口径", "包含需求深化、设计、开发、自测、联调、迁移、测试配合和上线支持"],
    ], [1.35, 5.15], fill=LIGHT_BLUE, font_size=9)

    callout(doc, "简要结论", f"本项目按当前修订口径约 {len(rows)} 个工作条目，正常工作量约 {normal_lo}-{normal_hi} 人日。SDD 对常规页面、脚手架、测试样例、文档和部分接口适配按 60% 以上提效目标组织；但核心业务、外部对接、报表口径、数据迁移和上线割接设置最低保护工时，因此全口径协同后工作量约 {vibe_lo}-{vibe_hi} 人日。该口径要求需求冻结、接口材料同步到位、自动化测试和人工评审强约束。")

    heading(doc, "一、简要评估", 1)
    add_table(doc, ["项目", "建议口径"], [
        ["功能规模", f"{len(rows)} 个功能项，覆盖原需求 14 个一级模块，并补充外部系统对接、数据湖接口和待定事项工作量。"],
        ["传统开发团队", "10-12 人，8-10 个月完成；包含产品、前端、后端、数据迁移、测试、实施。"],
        ["SDD 激进团队", "8-10 人，4-5 个月完成；前提是需求冻结、接口材料及时、测试自动化、代码评审、业务验收和并行联调节奏全部到位。"],
        ["SDD 稳妥团队", "8-10 人，5-6 个月完成；适用于接口材料存在补充、报表口径需要多轮确认、迁移数据质量需要治理的情况。"],
        ["外部条件受阻场景", "6 个月以上；适用于数据湖、IC卡、胆款、电子签章等接口资料或联调环境延期，或安全合规要求后置补充。"],
        ["核心难点", "票号段连续性、库存流水、票单结算、报表口径、数据权限、历史数据迁移、数据湖/外部系统对接、国产化适配。"],
        ["报价建议", "完整建设建议按 450-650 万区间，稳妥报价点 520-580 万；效率提升主要压缩周期和重复劳动，不等于按同等比例降低报价；电子签章、等保测评、第三方授权另列。"],
    ], [1.4, 5.1], fill=LIGHT_BLUE, font_size=9)

    heading(doc, "二、模块工作量简表", 1)
    add_table(doc, ["模块", "功能项数", "正常人日", "协同后人日", "重点说明"], module_summary(rows), [1.25, 0.65, 0.85, 0.9, 2.85], fill=LIGHT)

    heading(doc, "三、SDD 实施方式", 1)
    add_para(doc, "本方案中的 SDD（Spec-Driven Development，规范驱动开发）不是直接让工具一次性生成完整系统，而是把需求、设计、编码、测试、文档拆成小批量可验证任务。每个功能都按照“需求卡片 → 数据结构 → 接口 → 页面 → 测试 → 评审 → 合并”的节奏推进。本文 SDD 一律指本节定义的开发方法论，与软件工程语境中作为交付物的 Software Design Description（软件设计说明书）不同。")
    numbered(doc, [
        "需求卡片化：每个页面或业务动作形成独立任务，明确角色、数据范围、输入输出、状态和验收标准。",
        "先底座后业务：先完成认证、菜单、权限、字典、日志、导入导出、文件存储等平台能力。",
        "核心域手工建模：票据库存、票号段、结算、退票、调账等领域模型由架构师和业务人员先确认，再进入编码。",
        "批量生成常规代码：常规列表、表单、接口、DTO、Mapper、权限按钮、导入模板可以批量生成并统一 review。",
        "测试前置：每个核心业务先写状态流转和金额/张数计算测试，避免只实现页面不验证账务逻辑。",
        "小步交付：每个模块独立演示、独立验收、独立回归，不把问题积压到总体验收阶段。",
    ])
    heading(doc, "3.1 可压缩与不可压缩工作", 2)
    add_table(doc, ["类别", "可压缩内容", "不可压缩/低压缩内容", "管理要求"], [
        ["常规 CRUD 与页面", "列表、表单、详情、权限按钮、DTO、Mapper、基础接口、前端页面可批量生成。", "业务字段确认、页面原型确认、用户体验验收。", "适合按 60% 以上提效估算。"],
        ["核心票据业务", "单据骨架、状态枚举、接口样板、测试样例可提效。", "票号连续性、库存流水、结算口径、退票回写、调账规则。", "设置最低保护工时，必须人工评审。"],
        ["报表统计", "查询页面、导出模板、聚合 SQL 初稿可提效。", "指标口径、历史数据比对、金额/张数验收、旧系统样例核对。", "口径确认和验收时间不按编码效率压缩。"],
        ["外部系统对接", "适配层、Mock 服务、接口日志、重试框架可提效。", "接口文档等待、网络开通、联调排期、外部系统问题定位。", "拆分内部开发工时和外部等待周期。"],
        ["数据迁移", "映射脚本、校验脚本、异常报告生成可提效。", "历史脏数据处理、业务确认、割接窗口、抽样核验。", "至少三轮演练，不压缩关键验证。"],
        ["部署运维", "部署脚本、巡检脚本、监控配置模板可提效。", "生产网络、安全审查、备份恢复演练、回滚演练。", "纳入上线准入条件。"],
    ], [1.0, 1.8, 2.05, 1.65], fill=LIGHT_BLUE, font_size=8)
    heading(doc, "3.2 核心工作最低保护工时", 2)
    add_table(doc, ["工作类型", "最低保护工时建议", "原因"], [
        ["票单结算、重新结算", "单项不低于 8-12 人日", "涉及金额、张数、票号段、重新结算和历史追溯，不能只按页面复杂度估算。"],
        ["退票、调票、调账", "单项不低于 8-12 人日", "会影响库存回写、流水一致性和报表口径。"],
        ["数据湖、IC 卡接口", "单接口不低于 15-20 人日", "需字段确认、鉴权、安全、日志、联调、失败补偿和验收样例。"],
        ["胆款接口", "单接口不低于 15-18 人日", "以清分日报同步、补偿、对账为主，按中等接口复杂度预留。"],
        ["银行清分导入/自动对账", "不低于 12-18 人日", "需文件模板、重复导入控制、异常明细、人工复核和对账规则。"],
        ["核心报表", "单张不低于 4-8 人日", "报表开发本身可提效，但口径和数据核对需要保留时间。"],
        ["数据迁移与割接", "专项不低于 50-80 人日", "迁移脚本、三轮演练、异常治理、割接和回滚不可省略。"],
    ], [1.7, 1.2, 3.6], fill=LIGHT)
    add_para(doc, "下表给出 SDD 协同后的实施护栏，适用于详细功能项清单。它不是传统正常工作量，而是协同开发后仍需保留的区间口径；数据迁移一行的 30-55 人日为协同后专项口径，与上一表传统 50-80 人日并存不冲突。", size=9.5)
    add_table(doc, ["工作类别", "协同后人日护栏", "适用功能项", "说明"], [
        ["票单结算/重新结算", "8-12 人日", "票单结算、票单重新结算", "按保护工时修正表 19，不再直接套用统一压缩系数。"],
        ["退票/调票/调账", "8-12 人日", "票袋退票、票柜退票、票袋调票、调账管理", "涉及库存回写、流水一致性和报表回写。"],
        ["数据湖/IC 卡接口", "15-20 人日", "数据湖、IC 卡/一卡通", "含字段确认、鉴权、安全、日志、联调和失败补偿。"],
        ["胆款接口", "15-18 人日", "票款清分/胆款", "以清分日报同步、补偿、对账为主，按中等接口复杂度预留。"],
        ["银行清分/自动对账", "12-18 人日", "银行清分清点、无人售票自动对账", "含文件模板、重复导入控制、异常明细和人工复核。"],
        ["核心报表", "4-8 人日", "日结、月结、清分、收入、趋势等核心报表", "报表口径和旧系统样例核对不按编码效率压缩。"],
        ["数据迁移与割接", "30-55 人日", "Oracle 到达梦迁移、三轮演练、增量割接", "协同可提效脚本和报告生成，但数据治理和演练需保留周期。"],
    ], [1.4, 1.05, 2.0, 2.05], fill=LIGHT_BLUE, font_size=8)

    heading(doc, "四、实施步骤简表", 1)
    add_table(doc, ["阶段", "周期", "主要工作", "交付物"], [
        ["0. 需求冻结与切片", "1.5-2 周", "需求清单、业务流程、角色权限、报表口径、接口范围、迁移对象确认；同步拆成可并行任务卡。", "需求矩阵、原型、报表口径表、接口清单、迁移清单、任务切片清单。"],
        ["1. 平台底座", "2.5-3 周", "登录、用户、角色、菜单、数据权限、日志、字典、阈值、消息、文件导入导出。", "可登录可授权的基础系统。"],
        ["2. 基础与库存", "4-5 周", "基础资料、票库、票柜、审批、库存余额和流水。", "库存闭环 MVP。"],
        ["3. 票单核心", "5-6 周", "票袋、配票、结算、重新结算、退票、票号查询。", "票务室核心业务闭环。"],
        ["4. 扩展业务与报表", "4-5 周", "无人售、充值、特殊业务、集团/分公司/票务室报表。", "完整业务与统计能力。"],
        ["5. 迁移联调上线", "4-6 周", "数据迁移、接口联调、UAT、性能、安全整改、上线演练。", "生产上线版本与运维交接材料。"],
    ], [0.9, 0.8, 3.05, 1.75], fill=LIGHT)
    add_callout = callout
    add_callout(doc, "周期口径说明", "上表阶段存在并行关系，不应简单相加。平台底座完成 60%-70% 后，基础信息、库存、报表模板、接口 Mock、迁移脚本可并行推进。若甲方接口材料、报表口径、历史样例数据无法同步提供，周期应切换到 5-6 个月或 6 个月以上口径。")

    heading(doc, "五、关键设计细节", 1)
    heading(doc, "5.1 库存与票号设计", 2)
    bullets(doc, [
        "采用“库存余额表 + 库存流水表 + 业务单据表”三层模型，不以单一库存字段承载全部业务历史。",
        "票号段字段统一包含票价、票种、票组、起号、止号、张数、状态、库存主体、来源单据。",
        "所有出入库动作必须写流水，流水记录来源、目标、数量、金额、操作人、操作时间和业务原因。",
        "连续票号段需校验重叠、断号、越界、已核销、已退票、已结算等状态。",
        "重新结算不得覆盖原记录，应形成更正单据和差异流水，保证历史可追溯。",
    ])
    heading(doc, "5.2 权限与数据范围", 2)
    bullets(doc, [
        "权限拆分为菜单权限、按钮权限、接口权限、数据范围权限和导出权限。",
        "集团管理员可分配全局权限；集团操作员按授权查看部分分公司；分公司角色默认本分公司；票务室角色只看本票务室。",
        "后端查询必须统一注入数据范围条件，不能只依赖前端隐藏菜单。",
        "导出、打印、报表接口必须复用同一套数据权限规则。",
    ])
    heading(doc, "5.3 报表设计", 2)
    bullets(doc, [
        "每张报表上线前先确认指标定义、筛选条件、统计维度、取数来源、金额精度、导出格式。",
        "集团和分公司同类报表尽量合并，用数据权限和维度字段控制展示结果。",
        "高频报表建立日汇总/月汇总表，避免每次查询实时扫大量流水。",
        "报表验收必须使用旧系统样例数据做金额和张数比对。",
    ])
    heading(doc, "5.3.1 消息中心、阈值预警与打印模板", 3)
    add_table(doc, ["能力", "设计内容", "待确认事项"], [
        ["消息中心", "统一承载阈值预警、申请待审核、审核结果、导入失败、接口失败、结账提醒等站内消息。支持未读/已读、弹窗提醒、按角色过滤。", "是否需要短信、企业微信、邮件等外部通知；提醒频率和免打扰规则。"],
        ["阈值预警", "阈值管理配置指标、阈值类型、适用组织、提醒对象；达到阈值时进入消息中心并在用户打开系统时提醒。", "具体阈值指标、计算频率、是否需要定时任务预计算。"],
        ["审批待办", "印制、调拨、核销、调账、退票、销毁等流程统一生成待办消息，支持跳转到对应单据。", "各流程审批层级和超时提醒规则。"],
        ["打印模板", "票单配票、票单结算、票单查询整合配票单/结算单打印；票单打印菜单删除但打印能力保留。", "纸张规格、打印机环境、页眉页脚、签字栏、是否套打。"],
        ["Excel 导出", "报表、库存、流水、查询列表提供导出；大数据量导出采用异步任务和下载中心。", "导出字段、文件格式、最大导出量、是否水印。"],
    ], [1.0, 3.45, 2.05], fill=LIGHT_BLUE, font_size=8)
    heading(doc, "5.4 数据迁移", 2)
    bullets(doc, [
        "迁移对象包括 Oracle 表、序列、视图、函数、历史业务数据和必要附件。",
        "先做数据体检，输出重复票号、缺失组织、异常状态、金额不平、非法日期等问题清单。",
        "至少执行三轮迁移演练：开发、测试、准生产，每轮生成记录数、金额、票号覆盖和异常清单。",
        "上线割接需要冻结旧系统写入、执行增量迁移、业务抽样核验，并保留回滚方案。",
    ])
    heading(doc, "5.5 数据湖与外部系统对接", 2)
    add_para(doc, "重新审查需求文档后，外部系统对接不能只作为技术预留处理。数据湖、二维码系统、银行清分、票款清分/胆款系统、IC 卡系统、电子签章均会影响接口设计、数据模型、测试样例、部署网络和验收边界。")
    add_table(doc, ["对接对象", "需求依据", "设计内容", "工作量建议", "待确认事项"], [
        ["公交集团数据湖", "3.3.1 要求对外提供统一标准数据接口，用于数据湖获取票务系统数据。", "建设数据服务接口、字段字典、鉴权加密、调用日志、频率控制、失败重试、样例数据。", "18-35 人日", "数据范围、推送/拉取模式、接口协议、传输频率、加密算法、验收样例。"],
        ["二维码系统", "3.3.1 与数据湖并列；附录中二维码报表页面拟删除。", "确认页面删除后是否仍需保留二维码相关数据服务；如保留，需建立接口或数据集。", "12-28 人日", "页面删除是否等于业务下线；历史二维码数据是否迁移；是否还要对外供数。"],
        ["银行清分清点系统", "3.3.2 明确银行无网环境，通过线下报表、手动导入方式对接。", "设计导入模板、字段校验、重复导入控制、失败明细、导入日志、人工确认流程。", "15-32 人日", "银行文件格式、清分日报字段、异常处理、是否需要自动对账。"],
        ["票款清分/胆款系统", "3.3.3 由票务系统导入数据后同步胆款系统，拓展部分待明确。", "设计清分日报同步接口、对账字段、同步状态、失败重试、补偿任务和接口日志。", "15-32 人日", "同步方向、字段口径、触发时点、接口环境、失败重传和对账责任。"],
        ["IC 卡/一卡通系统", "3.3.4；线路、车辆、司售等基础信息来自 IC 卡系统同步。", "建设主数据同步、增量更新、异常补偿、数据映射、同步日志和人工修正入口。", "18-36 人日", "接口文档、同步频率、主数据归属、历史数据映射、异常数据处理。"],
        ["电子签章系统", "功能需求中说明根据业务后续评估再升级，特殊业务提供在线签字确认。", "一期预留签章对象、签章状态、附件归档和业务钩子；正式接口联调可单列。", "8-20 人日", "签章厂商、CA/证书、签章流程、法律效力、附件存储、费用授权。"],
    ], [1.05, 1.35, 1.75, 0.75, 1.6], fill=LIGHT_BLUE, font_size=7.8)
    heading(doc, "5.5.1 数据湖接口细化设计", 3)
    add_table(doc, ["设计项", "建议方案", "需要确认"], [
        ["数据目录", "建立票务主数据、库存流水、单据审批、票单结算、充值收入、无人售清分、报表汇总等数据集目录。", "数据湖实际需要哪些主题域，是否需要历史全量数据。"],
        ["同步模式", "支持全量初始化 + 增量同步；增量可按更新时间、业务流水号或同步批次号识别。", "数据湖是主动拉取还是票务系统主动推送。"],
        ["接口形态", "优先 REST/OpenAPI；大批量历史数据可采用文件交换或定时批量导出。", "接口协议、网络通道、文件格式、字符集和压缩方式。"],
        ["安全与加密", "接口鉴权、签名、时间戳、防重放、传输加密、敏感字段脱敏、调用白名单。", "集团安全要求、密钥管理方式、是否要求国密算法。"],
        ["数据质量", "提供记录数、金额合计、张数合计、批次状态、失败明细和重传机制。", "数据湖验收规则、质量阈值、错误处理责任边界。"],
        ["版本管理", "字段字典和接口版本化，新增字段向后兼容，重大变更需变更单。", "数据湖侧是否有统一接口规范和版本发布流程。"],
        ["审计监控", "记录调用方、接口、时间、参数摘要、返回状态、耗时、错误信息和重试次数。", "日志留存周期、审计字段、告警推送方式。"],
    ], [1.1, 3.25, 2.15], fill=LIGHT)
    heading(doc, "5.6 待定事项与范围冻结", 2)
    add_table(doc, ["待定事项", "需求文档表述", "影响", "冻结前处理建议"], [
        ["第三方对接模块", "待提供具体接口材料。", "影响接口工作量、排期、联调环境和报价边界。", "把接口文档、样例数据、联调账号、验收规则作为开工前置条件。"],
        ["安全性要求", "待集团补充确认。", "影响账号策略、日志留存、接口加密、部署网络和等保/密评。", "形成安全需求清单，明确是否含等保、密评、信创测评配合。"],
        ["调账管理", "具体业务场景待明确。", "影响库存、金额、报表和审计追溯。", "明确可调对象、审批层级、是否影响历史报表、是否需要会计凭证。"],
        ["特殊票管理", "具体业务场景待明确。", "影响是否纳入库存流水、收入统计和报表。", "明确特殊票类型、领取/核销/结算流程和数据口径。"],
        ["无人售票自动对账", "功能范围中提出新增无人售票自动对账、数据清洗。", "影响银行清分导入、异常处理和对账算法。", "明确银行文件格式、对账口径、容差规则、人工复核流程。"],
        ["电子签章", "后续根据业务评估再升级。", "影响特殊业务闭环和第三方费用。", "一期只做预留还是完整接入需在合同中明确。"],
    ], [1.2, 1.75, 1.65, 1.9], fill=LIGHT, font_size=8)
    heading(doc, "5.6.1 范围边界与报价口径", 3)
    add_table(doc, ["事项", "建议纳入一期", "建议单列/待确认", "原因"], [
        ["数据湖接口", "接口框架、字段字典、调用日志、Mock 联调、已确认数据集。", "数据湖侧未确认的数据集、特殊加密、历史全量大批量交换。", "数据范围和集团安全规范会直接影响工作量。"],
        ["二维码系统", "保留必要数据接口占位和历史数据确认。", "若页面删除后仍要求完整二维码统计供数，应单列确认。", "需求中页面拟删除，但 3.3.1 仍提到二维码系统。"],
        ["电子签章", "签章状态、附件字段、业务钩子、归档能力预留。", "完整电子签章厂商联调、CA 授权、法律效力流程。", "需求明确后续评估再升级。"],
        ["等保/密评/信创测评", "基础安全设计、日志审计、国产化适配。", "正式测评整改、测评材料、第三方平台对接。", "需求中安全性待集团补充确认。"],
        ["调账/特殊票", "基础单据与菜单预留。", "完整业务规则、报表影响、会计口径调整。", "原文标注具体业务场景待明确。"],
        ["无人售自动对账", "导入、清分日报、异常记录基础能力。", "自动对账算法、容差规则、异常闭环和清洗策略。", "原文提出新增能力但规则未展开。"],
    ], [1.05, 2.0, 2.0, 1.45], fill=LIGHT_BLUE, font_size=8)
    add_table(doc, ["序号", "冻结对象", "冻结物", "责任方", "截止节点", "未冻结后果"], [
        ["1", "二维码业务边界", "业务下线/仅删菜单/仍需对外供数三选一书面确认", "集团业务部", "第 1 周", "进入稳妥计划"],
        ["2", "调账管理业务场景", "可调对象、审批层级、是否影响历史报表、是否需会计凭证", "集团财务/业务", "第 2 周", "单列变更"],
        ["3", "特殊票管理业务场景", "特殊票类型、领取/核销/结算流程", "集团业务", "第 2 周", "单列变更"],
        ["4", "银行清分文件格式", "字段、分隔符、字符集、对账容差", "银行/集团清算", "第 3 周", "银行清分项延后"],
        ["5", "安全合规等级", "等保二级/三级、是否含密评/信创测评", "集团信息安全", "第 4 周", "安全整改后置或单列"],
        ["6", "电子签章范围", "一期预留/一期完整接入 + CA 厂商", "集团法务/业务", "第 8 周", "一期仅做预留"],
        ["7", "数据湖数据集与协议", "主题域、推/拉、加密算法、验收样例", "数据湖团队", "第 4 周", "数据湖仅交付框架"],
        ["8", "IC 卡同步主数据归属", "线路/车辆/司售数据源、增量字段", "一卡通公司", "第 4 周", "协同人日上调"],
        ["9", "报表口径样例", "集团/分公司/票务室各类报表旧系统样例数据", "集团业务", "第 6 周", "报表验收延后"],
        ["10", "达梦双套许可", "生产主备 + 开发测试主备", "集团采购", "第 6 周", "联调延期 1-2 周或本地用 H2/MySQL 临时联调"],
    ], [0.35, 1.05, 2.0, 1.0, 0.75, 1.35], fill=LIGHT, font_size=7.5)

    heading(doc, "六、部署与可运维性设计", 1)
    add_para(doc, "本项目属于票务核心管理系统，部署和运维能力应在一期同步设计，不能等到开发完成后补。运维设计目标是：可部署、可监控、可审计、可备份、可恢复、可灰度、可回滚。")
    heading(doc, "6.0 基础资源与容量规划", 2)
    add_table(doc, ["资源项", "原始需求口径", "设计建议"], [
        ["用户规模", "总量 >200 名，峰值 150 名在线。", "按 150 在线用户进行压力测试，报表、导入、结账等高负载场景单独压测。"],
        ["服务器资源", "3 台服务器：2 台数据库服务器，1 台应用服务器；总计 24vCPU、96G 内存。", "基础方案按 1 台应用主机交付；应用双节点仅作为建议位预留或二期可选，不纳入基础资源口径。"],
        ["数据库", "达梦数据库 2 套：生产 1 套主备 + 开发测试 1 套主备。", "达梦许可由甲方采购或在合同中单列；SQL 兼容性需前置验证，重点关注 Oracle 迁移函数、序列、视图和报表 SQL。"],
        ["历史数据", "Oracle 历史数据约 200GB，涉及 79 张表、序列、视图、函数等。", "先做数据体检和迁移映射，历史数据全量迁移后再做增量割接。"],
        ["对象存储", "10 年存储需求约 800-1000G OBS。", "导入文件、导出报表、签章附件、迁移中间文件需制定生命周期和归档策略。"],
        ["网络资源", "业务应用、数据库、数据中心服务器、线网中心内网需互通；涉及 192.168.99.x、10.129.87.x 和外网映射。", "上线前完成网络连通性、端口、白名单、接口访问、文件交换路径和外网 NAT 映射确认。"],
    ], [1.0, 2.25, 3.25], fill=LIGHT_BLUE, font_size=8)
    callout(doc, "资源与许可边界", "基础部署严格对齐原始需求 3 台服务器口径。应用双节点、双活容灾、第三方运维平台对接、OBS 异地副本属于建议项或二期可选。达梦数据库授权、国产服务器与操作系统、国产中间件与浏览器、电子签章 CA 与硬件介质、等保/密评/信创测评费用，由甲方采购或在合同中单列。")
    add_table(doc, ["运维领域", "设计要求", "工作量建议"], [
        ["环境规划", "至少规划开发、测试、预生产、生产四类环境；生产采用应用服务与数据库分离部署，数据库主备。", "12-18 人日"],
        ["国产化适配", "适配达梦数据库、国产服务器操作系统、国产中间件或 Nginx/应用服务部署要求，提前验证驱动和 SQL 兼容性。", "20-35 人日"],
        ["配置管理", "数据库连接、文件存储、外部接口、导入目录、日志级别、任务开关全部配置化，避免写死在代码中。", "8-12 人日"],
        ["CI/CD", "建立构建、单元测试、打包、制品归档、部署脚本流程；生产部署需支持版本标记和回滚。", "15-25 人日"],
        ["日志审计", "区分业务操作日志、登录日志、接口日志、导入导出日志、系统错误日志；关键单据状态变化必须可追溯。", "15-25 人日"],
        ["监控告警", "监控应用存活、接口耗时、数据库连接池、慢 SQL、磁盘空间、导入任务、定时任务和异常错误率。", "15-25 人日"],
        ["备份恢复", "数据库全量/增量备份，导入文件和导出附件备份，定期恢复演练，形成恢复时间目标。", "12-20 人日"],
        ["上线回滚", "上线前冻结窗口、迁移脚本、冒烟检查、业务抽样、失败回滚、旧系统只读策略均需预案。", "15-25 人日"],
        ["运维交接", "提供部署手册、配置手册、巡检手册、常见问题、数据恢复手册和接口联调手册。", "10-18 人日"],
        ["外部接口运维", "对数据湖、IC卡、胆款、银行清分导入等接口建立调用日志、失败告警、重试补偿和手工补传工具。", "20-35 人日"],
    ], [1.15, 4.2, 1.15], fill=LIGHT_BLUE)
    heading(doc, "6.1 运维验收指标建议", 2)
    add_table(doc, ["指标", "建议值/要求", "说明"], [
        ["RPO", "核心业务数据建议 ≤ 15 分钟；最终以甲方灾备要求和上线前验证为准。", "通过达梦主备同步和定时归档备份兜底；如要求 RPO ≤ 5 分钟需开启同步复制并单独验收。"],
        ["RTO", "普通故障 2 小时内恢复；重大故障按应急预案切换或回滚。", "需在上线前完成恢复演练和回滚演练。"],
        ["日志留存", "操作日志、登录日志、接口日志建议不少于 180 天；审计类日志可按集团要求延长。", "满足操作追溯、接口追踪和安全审计。"],
        ["接口重试", "外部接口失败建议自动重试 3 次，仍失败进入补偿队列并告警。", "适用于数据湖、IC 卡、胆款等接口。"],
        ["导入失败处理", "导入失败必须生成失败明细，可下载、可修正、可重新导入。", "适用于银行清分、无人售、残币、充值等文件导入场景。"],
        ["慢 SQL 阈值", "建议 2 秒以上记录慢 SQL，5 秒以上告警或进入优化清单。", "重点关注报表查询、库存查询和迁移后 SQL。"],
        ["备份恢复演练", "上线前至少 1 次，正式运行后建议每季度 1 次。", "演练结果纳入运维交接材料。"],
        ["生产配置管理", "密钥、数据库连接、接口地址、文件路径、任务开关全部外置配置。", "避免生产环境硬编码和手工改包。"],
    ], [1.0, 2.45, 3.05], fill=LIGHT)
    heading(doc, "6.2 性能、可靠性与兼容性验收", 2)
    add_table(doc, ["验收项", "建议指标", "说明"], [
        ["在线用户", "按峰值 150 在线用户设计压测场景。", "覆盖登录、列表查询、库存查询、票单结算、报表查询。"],
        ["普通查询", "常规列表查询建议 2 秒内返回。", "大数据量查询需分页、索引和权限过滤优化。"],
        ["报表查询", "常用日报/月报建议 5-10 秒内返回；复杂历史报表可异步生成。", "高频报表建议使用汇总表或预计算。"],
        ["Excel 导出", "大文件导出使用异步任务，完成后下载，避免浏览器长时间等待。", "需限制单次最大导出量和导出权限。"],
        ["文件导入", "导入任务异步执行，失败明细可下载，可修正后重导。", "适用于银行清分、无人售、残币、充值等模板。"],
        ["可靠性", "关键业务异常必须事务回滚，不产生半库存、半流水状态。", "票据库存、结算、退票、调账为重点。"],
        ["浏览器兼容", "优先支持 Chrome、Edge 及甲方内网指定国产浏览器。", "打印、上传、下载需在目标浏览器验证。"],
        ["打印兼容", "配票单、结算单需验证纸张、页边距、分页和打印机兼容。", "票单打印菜单删除，但打印能力保留到查询/结算入口。"],
    ], [1.05, 2.25, 3.2], fill=LIGHT)
    add_table(doc, ["压测场景", "建议并发/数据量", "通过判据", "工具建议"], [
        ["登录与菜单加载", "150 在线用户，分批登录", "登录成功率稳定，菜单加载无明显超时", "JMeter/Locust 或甲方指定工具"],
        ["库存查询", "50 并发，按分公司/票务室/票号段查询", "常规查询 2 秒内，大查询分页返回", "JMeter/Locust"],
        ["票单结算", "20 并发，含票号段计算和保存", "事务一致，无半库存/半流水状态", "接口压测 + 场景脚本"],
        ["报表查询", "20 并发，日报/月报/清分报表", "常用报表 5-10 秒内，复杂报表异步生成", "接口压测"],
        ["Excel 导出", "10 并发，大数据量导出", "进入异步任务，不阻塞浏览器请求", "接口压测 + 文件校验"],
        ["文件导入", "5 并发，银行清分/残币/充值模板", "失败明细准确，可修正后重导", "批量导入脚本"],
    ], [1.1, 1.4, 2.55, 1.45], fill=LIGHT_BLUE, font_size=8)
    add_para(doc, "压测完成后需执行数据一致性核验脚本，至少校验：库存余额等于流水汇总、票号段数量等于起号至止号枚举数量、单据金额等于明细金额汇总。核验结果应纳入性能测试报告。", size=9.5)
    heading(doc, "6.3 交付物、培训与运维交接", 2)
    add_table(doc, ["交付类别", "交付物", "说明"], [
        ["需求与设计", "需求矩阵、功能清单、原型、数据库设计、接口设计、报表口径表。", "作为验收范围和变更管理基线。"],
        ["开发交付", "源代码、构建脚本、部署包、数据库脚本、迁移脚本、导入模板、打印模板。", "版本需可追溯，生产发布包需归档。"],
        ["测试交付", "测试用例、接口测试报告、性能测试报告、UAT 问题清单、回归测试报告。", "重点覆盖库存流水、结算、权限、报表、迁移。"],
        ["数据迁移", "迁移方案、字段映射、迁移脚本、数据体检报告、迁移核对报告、割接记录。", "至少三轮迁移演练后上线。"],
        ["运维文档", "部署手册、配置手册、巡检手册、备份恢复手册、接口补偿手册、应急回滚预案。", "交付给甲方运维团队。"],
        ["用户材料", "用户操作手册、管理员手册、培训课件、常见问题清单。", "按集团、分公司、票务室角色分别组织培训。"],
        ["培训安排", "管理员培训、业务操作培训、报表核对培训、运维培训。", "建议至少 2 轮：上线前集中培训，上线后问题复盘培训。"],
    ], [1.05, 2.65, 2.8], fill=LIGHT_BLUE, font_size=8)
    heading(doc, "6.4 目标环境矩阵", 2)
    add_table(doc, ["维度", "必测项", "选测项", "备注"], [
        ["浏览器", "Chrome 100+、Edge 100+", "奇安信浏览器、UOS 浏览器、360 安全浏览器", "甲方信创浏览器需提供具体版本。"],
        ["操作系统", "Windows 10/11", "统信 UOS、麒麟 V10", "服务端和客户端分别声明。"],
        ["打印机", "1 款激光 + 1 款针打（甲方提供型号）", "热敏打印机（如有）", "票务打印多用针打，纸张需现场调试。"],
        ["屏幕", "1920×1080 桌面", "1366×768 笔记本、4K 屏", "UI 适配最小宽度建议 1366。"],
        ["网络", "内网 100M/1G", "跨区域内网带宽 ≥ 50M", "报表和导出场景按 50M 估算。"],
    ], [0.85, 1.55, 1.75, 2.35], fill=LIGHT, font_size=8)
    callout(doc, "部署工作量口径", "部署与运维建设建议单独预留 140-235 人日，其中外部接口运维、失败补偿和告警监控需纳入一期。如果甲方要求等保测评、密评配合、信创测评、双活容灾或第三方运维平台对接，应另行增加专项工作量和报价。")

    heading(doc, "七、详细功能项清单", 1)
    add_para(doc, "以下清单按需求文档附录逐项展开。正常人日为完整开发、自测、联调配合口径；协同后人日为采用 SDD 方式后的执行估算，不包含甲方等待、接口资料延迟和重大需求变更。")
    make_landscape(doc)
    heading(doc, "详细功能项与工作量估算", 1)
    detail_rows = []
    for r in rows:
        detail_rows.append([
            r["idx"],
            r["module"],
            r["page"],
            r["kind"],
            r["function"] or "按页面说明",
            r["dev"],
            f"{r['lo']}-{r['hi']}",
            r["vibe_display"],
            r["detail"],
        ])
    add_table(
        doc,
        ["序号", "模块", "功能页面", "类型", "现有功能", "开发内容", "正常人日", "协同后", "难点/备注"],
        detail_rows,
        [0.35, 0.95, 1.35, 0.75, 1.25, 2.1, 0.58, 0.58, 1.6],
        fill=LIGHT_BLUE,
        font_size=7.2,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "八、风险与边界", 1)
    add_table(doc, ["风险点", "表现", "建议处理"], [
        ["需求边界不清", "文档中存在“待明确”“删除”“整合”项，容易反复变更。", "在开发前形成冻结版功能清单，变更走审批和工期调整。"],
        ["接口资料缺失", "数据湖、二维码、胆款、电子签章等接口材料未完全明确。", "一期只做接口框架和已确认接口，未确认部分单独列二期。"],
        ["报表口径争议", "同一报表集团/分公司/票务室口径不同，数字验收困难。", "先用旧系统样例数据签字确认，再开发。"],
        ["历史数据异常", "旧系统长期运行导致脏数据和状态不一致。", "迁移前做数据体检，异常数据由业务确认处理规则。"],
        ["国产化兼容风险", "达梦 SQL、数据库驱动、国产操作系统、中间件和浏览器兼容问题可能后置暴露。", "开发早期即接入达梦和目标浏览器，迁移函数、序列、视图和报表 SQL 前置验证。"],
        ["网络连通延期", "云平台、数据中心、线网中心、外部接口白名单和端口开通不及时。", "把网络连通、端口、白名单和 NAT 映射列为上线前置条件。"],
        ["安全合规后补", "等保、密评、信创测评要求若后置补充，会影响架构、日志、加密和部署。", "安全等级和测评范围在第 4 周前冻结，正式测评配合单列报价。"],
        ["打印与浏览器兼容", "配票单、结算单在目标浏览器、打印机、纸张和页边距下可能存在偏差。", "上线前用甲方实际浏览器和打印机做专项验证。"],
        ["旧系统并行窗口", "上线割接后旧系统是否只读、并行多久、如何回滚未明确。", "制定割接窗口、只读策略、增量迁移、回滚分支和抽样验收流程。"],
        ["AI 生成代码质量波动", "常规代码效率高，但核心业务可能遗漏隐性规则。", "核心交易类代码必须人工 review，并要求测试覆盖状态流转和金额计算。"],
    ], [1.35, 2.45, 2.7], fill=LIGHT)

    heading(doc, "九、结论", 1)
    add_para(doc, f"本项目按当前修订口径共 {len(rows)} 个工作条目，已将原需求明确删除/整合的功能按 0-1 人日处理，并合并集团/分公司同名报表。传统完整实施工作量约 {normal_lo}-{normal_hi} 人日，SDD 协同后约 {vibe_lo}-{vibe_hi} 人日；核心业务、外部接口、报表口径、数据迁移和上线割接设置最低保护工时，不随 60% 提效目标等比例压缩。")
    add_para(doc, "周期建议采用双口径：激进计划为 8-10 人、4-5 个月，前提是需求冻结、接口资料、历史样例数据、达梦许可和联调环境在约定节点全部到位；稳妥计划为 8-10 人、5-6 个月，适用于接口材料补充、报表口径多轮确认或迁移数据质量治理场景。任一关键前置条件未达成，应自动切换至稳妥计划或单列变更。上述协同后人日不含部署与运维建设的 140-235 人日；后者已包含在基础报价中，仅当甲方追加等保/密评/信创测评、双活容灾或第三方运维平台对接时另行增加专项报价。")
    add_para(doc, "报价建议维持 450-650 万区间，稳妥报价点 520-580 万。该报价不含达梦数据库授权、国产服务器与操作系统、国产中间件、第三方电子签章、CA/硬件介质、等保/密评/信创测评、第三方运维平台和双活容灾等单列项目。效率提升主要压缩周期和重复劳动，不等于按同等比例降低项目责任、质保、联调和验收成本。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
