from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SRC = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/北京公交集团票务综合管理平台系统升级-SDD实施方案与工作量清单.docx")
OUT = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/北京公交集团票务综合管理平台系统升级-SDD实施方案与工作量清单-含设计图.docx")
PNG = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/svg设计图/png预览")


DIAGRAMS = {
    "三、SDD 实施方式": ("05-SDD协同交付流程图.png", "图 1 SDD 协同交付流程图"),
    "四、实施步骤简表": ("07-系统功能地图.png", "图 2 系统功能地图"),
    "五、关键设计细节": ("02-核心票务业务流转图.png", "图 3 核心票务业务流转图"),
    "5.2 权限与数据范围": ("04-权限与数据范围设计图.png", "图 4 权限与数据范围设计图"),
    "六、部署与可运维性设计": ("03-部署与可运维性架构图.png", "图 5 部署与可运维性架构图"),
    "5.5 数据湖与外部系统对接": ("06-数据湖与外部系统对接架构图.png", "图 6 数据湖与外部系统对接架构图"),
    "七、详细功能项清单": ("01-总体技术架构图.png", "图 7 总体技术架构图"),
}


def insert_after(paragraph, image_path: Path, caption: str, width=Inches(6.7)):
    caption_p = paragraph.insert_paragraph_before("")
    # Move caption and picture after the heading by inserting before next paragraph is not directly
    # supported in python-docx. Use low-level insertion after heading.
    pic_p = paragraph._p.addnext(paragraph._p.__class__())


def add_picture_after(paragraph, image_path: Path, caption: str, width=Inches(6.7)):
    doc = paragraph.part.document
    p_pic = doc.add_paragraph()
    p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pic.add_run()
    run.add_picture(str(image_path), width=width)
    p_pic.paragraph_format.space_before = Pt(6)
    p_pic.paragraph_format.space_after = Pt(3)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(8)
    r = p_cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.name = "Microsoft YaHei"

    # Move the newly appended picture + caption immediately after the target paragraph.
    paragraph._p.addnext(p_cap._p)
    paragraph._p.addnext(p_pic._p)


def main():
    doc = Document(SRC)
    inserted = 0
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text in DIAGRAMS:
            file_name, caption = DIAGRAMS[text]
            image = PNG / file_name
            if not image.exists():
                raise FileNotFoundError(image)
            width = Inches(8.9) if "详细功能项" in text else Inches(6.7)
            add_picture_after(para, image, caption, width=width)
            inserted += 1
    if inserted != len(DIAGRAMS):
        raise RuntimeError(f"Inserted {inserted}, expected {len(DIAGRAMS)}")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
