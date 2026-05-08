from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/北京公交集团票务综合管理平台系统升级-详细设计与Codex实施方案.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(96, 108, 122)
LIGHT = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", style: str | None = None, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=LIGHT) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, size=9.5, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for r in p.runs:
                    set_run_font(r, size=9)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    set_table_width(table, [6.5])
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "北京公交集团票务综合管理平台系统升级 | 详细设计与 Codex 实施方案"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in header.runs:
        set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "内部评估文件"
    for r in footer.runs:
        set_run_font(r, size=9, color=MUTED)


def build() -> None:
    doc = Document()
    configure_document(doc)

    add_para(doc, "技术方案", bold=True, color=MUTED, size=11)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("北京公交集团票务综合管理平台系统升级")
    set_run_font(r, size=24, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("详细设计与 Codex 协作实施方案")
    set_run_font(r, size=15, color=MUTED)

    meta = [
        ["文档用途", "用于项目立项、报价拆分、研发实施和 Codex 协作开发规划"],
        ["编制日期", date.today().isoformat()],
        ["需求来源", "《北京公交集团票务综合管理平台系统升级-需求分析说明书（初稿0506）》"],
        ["版本", "V1.0"],
    ]
    add_table(doc, ["字段", "内容"], meta, [1.3, 5.2], header_fill=LIGHT_BLUE)

    add_callout(
        doc,
        "总体判断",
        "本项目属于存量核心业务系统重构升级，重点不是页面生成，而是票据库存流水、票单结算、报表口径、数据权限、历史数据迁移和国产化适配。Codex 可显著提升需求整理、重复代码生成、测试补齐和文档产出效率，但架构边界、业务规则确认和上线割接仍需人工主导。",
    )

    add_heading(doc, "1. 项目范围与建设目标", 1)
    add_para(doc, "本次建设目标是在继承现有生产系统主要能力的基础上，完成业务流程重构、国产化数据库迁移、权限和数据范围治理、报表整合、接口标准化和运维能力提升。系统面向集团、分公司和票务室多级用户，覆盖有人售票、无人售票、票库、票柜、票袋、票单、充值卡、特殊业务和统计报表等票务经营管理场景。")
    add_bullets(doc, [
        "完整保留并优化核心票务业务闭环：印制、调拨、核销、配票、结算、退票、库存查询和流水追溯。",
        "整合同类报表，按集团、分公司、票务室三级数据权限展示统计结果。",
        "完成 Oracle 历史数据到达梦数据库的迁移、校验和兼容适配。",
        "预留并分阶段实现与 IC 卡系统、数据湖、二维码系统、票款清分系统、电子签章系统的对接。",
        "建立可审计、可测试、可持续迭代的工程体系，支持 Codex 辅助开发和自动化校验。",
    ])

    add_heading(doc, "2. 总体技术架构", 1)
    add_table(doc, ["层级", "建议技术", "设计说明"], [
        ["前端展示层", "Vue 3 + TypeScript + Vite + Element Plus", "承载菜单导航、列表查询、业务表单、报表展示、打印和导入导出入口。"],
        ["后端服务层", "Java 17 + Spring Boot 3 + MyBatis-Plus", "实现业务编排、事务控制、权限校验、库存流水、报表聚合和接口服务。"],
        ["数据存储层", "达梦 DM8 + Redis + OBS/MinIO", "达梦作为主业务库，Redis 用于会话/缓存/短期任务状态，文件对象存储用于模板、导入文件和导出文件。"],
        ["集成层", "REST/OpenAPI + 批量文件导入 + 定时任务", "对外提供标准接口；对无网络银行清分场景保留线下 Excel/CSV 导入；定时同步基础信息和报表数据。"],
        ["运维层", "Nginx + 应用服务 + 日志审计 + 监控告警", "满足国产化部署、日志追溯、性能监控和上线回滚要求。"],
    ], [1.15, 1.75, 3.6], header_fill=LIGHT_BLUE)

    add_heading(doc, "3. 系统模块设计", 1)
    modules = [
        ["基础信息", "票务室、线路、票价、车辆、司售、普票类型、充值类型、机具、多样化运营类型、IC 卡类型", "以 CRUD 和同步查询为主，需统一数据权限和停用机制。"],
        ["票库管理", "印制、调拨、核销、调账、票库库存、票库出入库查询", "核心在审批状态、库存流水、票号段合法性和跨公司调拨。"],
        ["票柜管理", "票柜库存、入库、出库、配票申请、票务室调票", "承接票库到票务室的库存流转，需要单张/批量调票能力。"],
        ["票单管理", "票袋、票单配票、票单结算、票单查询、票袋调票、重新结算、票号查询", "项目最复杂模块，需要严格状态机和结算口径。"],
        ["无人售与特殊业务", "异物、大额币、残币、无人售线路结算、多样化收入、特殊票", "重点是异常状态、待销毁/领取/入账流程和批量导入。"],
        ["退票管理", "票袋退票、票柜退票", "支持连续退票和单张退票，必须回写库存和流水。"],
        ["充值管理", "IC 卡库存、入库、出库、网点库存、售出、充值收入", "需要处理网点维度库存、售出导入和收入记录。"],
        ["统计报表", "集团、分公司、票务室多级报表、趋势图、Excel 导出", "需先冻结报表口径，避免开发完成后重复返工。"],
        ["系统管理", "用户、角色、权限、菜单、阈值、日志", "为全系统提供认证、授权、数据范围和审计能力。"],
    ]
    add_table(doc, ["模块", "主要功能", "设计重点"], modules, [1.25, 2.9, 2.35])

    add_heading(doc, "4. 核心业务状态与流水设计", 1)
    add_para(doc, "票据类业务必须采用“当前库存 + 流水台账 + 状态机”的组合模型。只存库存余额无法满足审计、调账、退票、票号查询和异常追溯要求。")
    add_table(doc, ["业务对象", "关键状态", "关键校验"], [
        ["印制申请", "待审核、不同意、待确认、完成", "待审核/不同意可修改；完成后生成票库入库流水。"],
        ["调拨申请", "待审核、不同意、待确认、完成", "校验调出库存、调入公司、票号段不重叠。"],
        ["核销申请", "待审核、完成、驳回", "核销后票号段不可再配出，保留原因和审批记录。"],
        ["配票申请", "待处理、已配票、退回、完成", "票柜库存扣减，票袋或票单库存增加。"],
        ["票单结算", "已配票、已结算、重新结算中、已更正", "销售张数和金额由剩余票号段反推，保留旧结算记录。"],
        ["退票", "申请、确认、完成、驳回", "退票票号不能已核销、不能重复退票，完成后回写对应库存。"],
    ], [1.3, 2.1, 3.1])

    add_heading(doc, "5. 数据模型设计原则", 1)
    add_bullets(doc, [
        "基础主数据表：组织、分公司、票务室、线路、车辆、司售、票价、票种、充值类型、机具等。",
        "库存余额表：按库存主体、票种、票价、票组、起止号、状态聚合当前可用数量。",
        "库存流水表：记录每一次入库、出库、调拨、退票、核销、调账、配票、结算的来源单据和操作人。",
        "业务单据表：每类申请、审核、确认、结算均独立建单，避免把审批状态混入库存余额。",
        "报表宽表/汇总表：对日结、月结、趋势报表建立可重算的汇总层，减少实时聚合压力。",
        "审计表：记录登录、菜单访问、增删改、导入导出、审批、结账等关键操作。",
    ])

    add_heading(doc, "6. 权限与数据范围设计", 1)
    add_table(doc, ["角色", "功能权限", "数据范围"], [
        ["集团管理员", "最高权限，可维护菜单、角色、用户和全局参数", "全集团数据，可分配集团/分公司/票务室数据权限。"],
        ["集团操作员", "按角色授权使用业务功能", "可查看被授权的分公司和票务室数据。"],
        ["分公司管理员", "维护本分公司及下属票务室用户和权限", "默认本分公司数据，可细分到票务室。"],
        ["分公司操作员", "使用分公司层面业务和报表功能", "本分公司及授权票务室数据。"],
        ["票务室操作员", "操作票柜、票袋、票单、结算等票务室业务", "仅本票务室数据。"],
    ], [1.25, 2.55, 2.7], header_fill=LIGHT_BLUE)
    add_callout(doc, "权限实现建议", "权限必须拆成菜单权限、按钮权限、数据范围权限和字段/导出权限。所有查询接口必须统一经过数据范围拦截器，避免前端隐藏菜单但后端接口越权。")

    add_heading(doc, "7. 接口与集成设计", 1)
    add_table(doc, ["系统/对象", "对接方式", "一期处理建议"], [
        ["IC 卡系统", "接口同步或定时任务", "优先同步线路、车辆、司售、IC 卡基础信息；接口不稳定时保留手动导入兜底。"],
        ["数据湖/二维码系统", "标准 REST/API 或数据推送", "先定义统一数据服务接口、鉴权、频率、字段口径；具体联调作为独立里程碑。"],
        ["银行清分清点系统", "线下 Excel/CSV 导入", "按银行无网环境处理，建立模板、校验、导入日志和失败明细下载。"],
        ["票款清分/胆款系统", "批量同步或接口推送", "以票务系统生成的清分日报为主数据源，明确对账字段。"],
        ["电子签章系统", "第三方接口", "建议一期预留签章对象、签章状态和附件字段，实际签章联调单独报价/排期。"],
    ], [1.55, 1.55, 3.4])

    add_heading(doc, "8. 数据迁移方案", 1)
    add_numbered(doc, [
        "盘点 Oracle 现有 79 张表、序列、视图、函数、触发器和历史附件，建立字段映射表。",
        "完成达梦库结构设计，明确字段类型、索引、主键、唯一约束和历史兼容字段。",
        "开发迁移脚本，先迁移基础主数据，再迁移库存、单据、结算、报表历史数据。",
        "执行数据清洗：重复票号、缺失组织、异常状态、金额不平、非法日期等问题生成清洗报告。",
        "进行三轮迁移演练：开发环境、测试环境、准生产环境，每轮输出数量核对和金额核对报告。",
        "上线割接时冻结旧系统写入，完成增量迁移、业务抽样验证和回滚预案确认。",
    ])

    add_heading(doc, "9. Codex 协作开发流程", 1)
    add_table(doc, ["阶段", "Codex 可承担工作", "人工控制点"], [
        ["需求结构化", "从 Word/Excel 提取功能项、生成需求池、补充初版验收点", "业务人员确认字段、流程、报表口径和删除/保留范围。"],
        ["设计阶段", "生成表结构草案、接口草案、状态机草案、测试场景清单", "架构师确认边界、事务、权限模型和关键业务规则。"],
        ["编码阶段", "生成 CRUD、Service、Mapper、Controller、前端页面、导入导出、测试用例", "开发负责人审查代码风格、事务完整性和异常处理。"],
        ["测试阶段", "补接口测试、构造 Mock 数据、生成回归脚本、修复构建问题", "测试负责人确认场景覆盖和真实业务数据抽样。"],
        ["交付阶段", "生成部署文档、接口文档、用户手册和变更说明", "项目组确认上线窗口、培训材料、应急预案。"],
    ], [1.2, 2.8, 2.5], header_fill=LIGHT_BLUE)

    add_heading(doc, "10. 实施计划", 1)
    add_table(doc, ["阶段", "周期", "主要产出"], [
        ["阶段 0：需求深化", "3-5 周", "最终需求清单、原型、报表口径、接口清单、迁移清单。"],
        ["阶段 1：平台底座", "4-6 周", "登录、权限、菜单、日志、字典、导入导出、消息中心。"],
        ["阶段 2：基础与库存", "8-10 周", "基础资料、票库、票柜、库存流水、审批流程。"],
        ["阶段 3：票单核心", "8-12 周", "票袋、配票、结算、重新结算、票号查询、退票。"],
        ["阶段 4：扩展业务与报表", "8-10 周", "无人售、充值、特殊业务、集团/分公司/票务室报表。"],
        ["阶段 5：迁移联调与上线", "8-10 周", "数据迁移、接口联调、UAT、性能测试、安全整改、上线割接。"],
    ], [1.55, 1.15, 3.8])

    add_heading(doc, "11. 测试与验收策略", 1)
    add_bullets(doc, [
        "单元测试：覆盖票号段计算、金额计算、库存扣减、状态流转、权限过滤等纯业务逻辑。",
        "接口测试：覆盖所有新增、修改、删除、审核、确认、结账、导入导出接口。",
        "场景测试：按印制到配票、配票到结算、结算到报表、退票回库等业务闭环执行。",
        "数据迁移测试：按表数量、记录数、金额合计、票号段覆盖、异常数据清单进行核对。",
        "权限测试：集团、分公司、票务室分别验证菜单、按钮、数据、导出范围。",
        "性能测试：以峰值 150 在线用户为基准，重点压测报表查询、导入、结账和库存查询。",
    ])

    add_heading(doc, "12. 风险与控制措施", 1)
    add_table(doc, ["风险", "影响", "控制措施"], [
        ["报表口径未冻结", "开发完成后反复返工", "每张报表先签字确认指标、维度、过滤条件和样例数据。"],
        ["历史数据质量差", "迁移延期、上线后账不平", "提前做数据体检，问题数据单独出清洗规则和责任确认。"],
        ["接口资料不完整", "联调延期", "接口作为独立里程碑，未提供资料的内容不阻塞核心业务开发。"],
        ["票据状态规则遗漏", "库存错误、审计风险", "核心单据全部状态机化，关键状态变更必须有流水和测试用例。"],
        ["Codex 生成代码缺少业务约束", "隐藏缺陷", "Codex 只产出初稿，核心业务必须人工 review 和场景测试。"],
        ["国产化兼容问题", "部署和性能不稳定", "早期引入达梦、国产 OS、中间件环境，不到上线前再适配。"],
    ], [1.65, 1.8, 3.05])

    add_heading(doc, "13. 结论建议", 1)
    add_para(doc, "建议采用“核心闭环优先、接口分阶段、报表先定口径、迁移提前演练”的实施策略。Codex 应作为研发加速工具嵌入需求整理、代码生成、测试补齐和文档交付过程，但不能替代架构决策、业务确认和验收责任。")
    add_bullets(doc, [
        "一期优先交付票务核心闭环：基础信息、票库、票柜、票袋、票单、退票、无人售、充值、报表、权限和迁移。",
        "电子签章、复杂实时接口、等保/密评配合建议单独列项，降低一期范围失控风险。",
        "开发组织建议按 10-12 人、8-10 个月主计划推进，使用 Codex 后可将重复开发和文档测试工作压缩 30%-50%。",
        "正式实施前应先输出功能清单 Excel、原型稿、数据库初设和迁移体检报告，作为报价和排期基线。",
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
